"""供 Worker 使用的最小工具集。所有文件操作均限制在项目目录内。"""

from __future__ import annotations

import os
import subprocess
import sys
from html import escape
from pathlib import Path
from typing import Any, Callable

ROOT_DIR = Path(__file__).resolve().parent


def _safe_path(path: str) -> Path:
    """解析相对路径，并阻止访问项目目录之外的文件。"""
    target = (ROOT_DIR / path).resolve()
    if target != ROOT_DIR and ROOT_DIR not in target.parents:
        raise ValueError("Path must stay inside the project directory")
    return target


def read_file(path: str) -> str:
    return _safe_path(path).read_text(encoding="utf-8")


def write_file(path: str, content: str) -> str:
    target = _safe_path(path)
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_text(content, encoding="utf-8")
    return f"Wrote {len(content)} characters to {path}"


def create_artifact(
    format: str,
    path: str,
    title: str,
    content: str = "",
    sheets: list[dict[str, Any]] | None = None,
) -> str:
    """Generate a Markdown, HTML, Word, or Excel artifact inside the project."""
    artifact_format = format.lower()
    extensions = {"markdown": ".md", "html": ".html", "word": ".docx", "excel": ".xlsx"}
    if artifact_format not in extensions:
        raise ValueError(f"Unsupported artifact format: {format}")
    target = _safe_path(path)
    if target.suffix.lower() != extensions[artifact_format]:
        raise ValueError(f"{artifact_format} artifacts must use {extensions[artifact_format]}")
    target.parent.mkdir(parents=True, exist_ok=True)

    if artifact_format == "markdown":
        target.write_text(f"# {title}\n\n{content.rstrip()}\n", encoding="utf-8")
    elif artifact_format == "html":
        paragraphs = "\n".join(
            f"<p>{escape(paragraph)}</p>" for paragraph in content.split("\n\n") if paragraph.strip()
        )
        document = (
            "<!doctype html>\n<html lang=\"en\"><head><meta charset=\"utf-8\">"
            f"<title>{escape(title)}</title></head><body><main><h1>{escape(title)}</h1>"
            f"{paragraphs}</main></body></html>\n"
        )
        target.write_text(document, encoding="utf-8")
    elif artifact_format == "word":
        from docx import Document

        document = Document()
        document.add_heading(title, level=0)
        for paragraph in content.split("\n\n"):
            if paragraph.strip():
                document.add_paragraph(paragraph.strip())
        document.save(target)
    else:
        from openpyxl import Workbook

        if not sheets:
            raise ValueError("Excel artifacts require at least one sheet")
        workbook = Workbook()
        workbook.remove(workbook.active)
        for sheet_data in sheets:
            worksheet = workbook.create_sheet(str(sheet_data.get("name") or "Sheet")[:31])
            for row in sheet_data.get("rows", []):
                worksheet.append(list(row))
        workbook.save(target)
    return f"Created {artifact_format} artifact at {path}"


def run_python(code: str) -> str:
    """执行短 Python 代码。生产环境应替换为容器或沙箱。"""
    if os.getenv("ENABLE_PYTHON_TOOL", "false").lower() != "true":
        return "Python execution is disabled. Set ENABLE_PYTHON_TOOL=true to enable it."

    timeout = int(os.getenv("PYTHON_TIMEOUT_SECONDS", "10"))
    try:
        result = subprocess.run(
            [sys.executable, "-c", code],
            cwd=ROOT_DIR,
            capture_output=True,
            text=True,
            timeout=timeout,
            check=False,
        )
    except subprocess.TimeoutExpired:
        return f"Python execution timed out after {timeout} seconds"

    output = result.stdout
    if result.stderr:
        output += f"\nSTDERR:\n{result.stderr}"
    return f"Exit code: {result.returncode}\n{output}".strip()


TOOL_FUNCTIONS: dict[str, Callable[..., str]] = {
    "read_file": read_file,
    "write_file": write_file,
    "create_artifact": create_artifact,
    "run_python": run_python,
}

TOOL_SCHEMAS: list[dict[str, Any]] = [
    {
        "type": "function",
        "function": {
            "name": "create_artifact",
            "description": "Create a Markdown (.md), HTML (.html), Word (.docx), or Excel (.xlsx) file inside the project directory.",
            "parameters": {
                "type": "object",
                "properties": {
                    "format": {"type": "string", "enum": ["markdown", "html", "word", "excel"]},
                    "path": {"type": "string", "description": "Relative output path with the matching extension."},
                    "title": {"type": "string"},
                    "content": {"type": "string", "description": "Body text for Markdown, HTML, and Word."},
                    "sheets": {
                        "type": "array",
                        "description": "Required for Excel. Each sheet contains a name and a 2D rows array.",
                        "items": {
                            "type": "object",
                            "properties": {
                                "name": {"type": "string"},
                                "rows": {"type": "array", "items": {"type": "array", "items": {}}},
                            },
                            "required": ["name", "rows"],
                        },
                    },
                },
                "required": ["format", "path", "title"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "read_file",
            "description": "Read a UTF-8 text file inside the project directory.",
            "parameters": {
                "type": "object",
                "properties": {"path": {"type": "string"}},
                "required": ["path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "write_file",
            "description": "Write a UTF-8 text file inside the project directory.",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string"},
                    "content": {"type": "string"},
                },
                "required": ["path", "content"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "run_python",
            "description": "Run a short Python snippet and return stdout/stderr.",
            "parameters": {
                "type": "object",
                "properties": {"code": {"type": "string"}},
                "required": ["code"],
            },
        },
    },
]


def call_tool(name: str, arguments: dict[str, Any]) -> str:
    function = TOOL_FUNCTIONS.get(name)
    if not function:
        return f"Unknown tool: {name}"
    try:
        return function(**arguments)
    except Exception as exc:  # 将工具错误反馈给模型，而不是中断整个任务
        return f"Tool error: {type(exc).__name__}: {exc}"
